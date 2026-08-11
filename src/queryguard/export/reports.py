"""Small report exporters used by the Streamlit UI."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from queryguard.models import DocumentQueryResponse


def document_answer_markdown(response: DocumentQueryResponse) -> str:
    lines = [
        "# QueryGuard Document Analysis",
        "",
        f"**Question:** {response.question}",
        "",
        "## Answer",
        "",
        response.answer or "No answer was generated.",
        "",
        "## Evidence",
        "",
    ]
    for index, source in enumerate(response.sources, start=1):
        lines.extend(
            [
                f"### Source {index}: {source.source_name} — {source.locator}",
                "",
                source.excerpt,
                "",
            ]
        )
    return "\n".join(lines)


def document_answer_docx_bytes(response: DocumentQueryResponse) -> bytes:
    document = Document()
    document.add_heading("QueryGuard Document Analysis", level=1)
    document.add_paragraph(f"Question: {response.question}")
    document.add_heading("Answer", level=2)
    document.add_paragraph(response.answer or "No answer was generated.")
    document.add_heading("Evidence", level=2)
    for index, source in enumerate(response.sources, start=1):
        document.add_heading(
            f"Source {index}: {source.source_name} — {source.locator}",
            level=3,
        )
        document.add_paragraph(source.excerpt)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
