"""DOCX text and table extraction with section-aware provenance."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from queryguard.documents.models import DocumentUnit
from queryguard.ingestion.common import (
    IngestionError,
    validate_extension,
    validate_office_archive,
)

DOCX_EXTENSIONS = {".docx"}


def extract_docx_units(
    path: Path,
    *,
    max_uncompressed_bytes: int,
) -> tuple[list[DocumentUnit], list[str]]:
    validate_extension(path, DOCX_EXTENSIONS)
    validate_office_archive(path, max_uncompressed_bytes=max_uncompressed_bytes)

    try:
        document = Document(path)
    except Exception as exc:
        raise IngestionError(f"Could not open DOCX {path.name}: {exc}") from exc

    units: list[DocumentUnit] = []
    section = "Document"
    paragraph_number = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        paragraph_number += 1
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            section = text
            continue
        units.append(
            DocumentUnit(
                source_name=path.name,
                locator=f"{section} · paragraph {paragraph_number}",
                text=text,
            )
        )

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            units.append(
                DocumentUnit(
                    source_name=path.name,
                    locator=f"Table {table_index}",
                    text="\n".join(rows),
                )
            )

    if not units:
        raise IngestionError(f"No readable text was found in {path.name}.")
    return units, []
